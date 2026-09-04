import streamlit as st
from google import genai
from google.genai import types
import pdfplumber
from docx import Document
from PIL import Image
import json
import re
import io

# ---------- PAGE CONFIG ----------
st.set_page_config(page_title="ATS Resume Scorer", page_icon="📄", layout="wide")

# ---------- SIDEBAR: API KEY ----------
st.sidebar.title("⚙️ Settings")

default_key = ""
try:
    default_key = st.secrets.get("GEMINI_API_KEY", "")
except Exception:
    default_key = ""

api_key = st.sidebar.text_input(
    "Enter your Gemini API Key", type="password", value=default_key
)
st.sidebar.markdown("[Get a free API key](https://aistudio.google.com/app/apikey)")

MODEL_NAME = "gemini-2.5-flash"  # alias: always resolves to Google's newest Flash model

client = None
if api_key:
    client = genai.Client(api_key=api_key)

# ---------- TEXT EXTRACTION FUNCTIONS ----------

def extract_text_from_pdf(file_bytes):
    text = ""
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        st.error(f"Error reading PDF: {e}")
    return text.strip()


def extract_text_from_docx(file_bytes):
    text = ""
    try:
        doc = Document(io.BytesIO(file_bytes))
        for para in doc.paragraphs:
            text += para.text + "\n"
        # also grab text inside tables (many resumes use tables)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
    except Exception as e:
        st.error(f"Error reading DOCX: {e}")
    return text.strip()


def extract_text_from_image(file_bytes, client):
    """Use Gemini's vision capability to read text directly from an image resume."""
    try:
        image = Image.open(io.BytesIO(file_bytes))
        response = client.models.generate_content(
            model=MODEL_NAME,
            contents=[
                "Extract ALL text from this resume image exactly as it appears, "
                "preserving section order (name, contact, summary, experience, "
                "education, skills, etc). Return ONLY the extracted text, no commentary.",
                image,
            ],
        )
        return response.text.strip()
    except Exception as e:
        st.error(f"Error reading image with Gemini Vision: {e}")
        return ""


def get_resume_text(uploaded_file, client):
    file_bytes = uploaded_file.read()
    file_type = uploaded_file.type

    if file_type == "application/pdf":
        text = extract_text_from_pdf(file_bytes)
        if not text:
            st.warning(
                "No selectable text found in PDF — it may be a scanned image. "
                "Try uploading it as an image instead for best results."
            )
        return text

    elif file_type in [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ]:
        return extract_text_from_docx(file_bytes)

    elif file_type in ["image/png", "image/jpeg", "image/jpg"]:
        return extract_text_from_image(file_bytes, client)

    else:
        st.error("Unsupported file type.")
        return ""


# ---------- GEMINI ATS ANALYSIS ----------

ATS_PROMPT_TEMPLATE = """
You are an expert ATS (Applicant Tracking System) analyzer and professional resume coach.

Analyze the following resume text and evaluate it as an ATS system would.

RESUME TEXT:
\"\"\"
{resume_text}
\"\"\"

{jd_section}

Return your analysis STRICTLY as a valid JSON object with this exact structure (no markdown fences, no extra text):

{{
  "ats_score": <integer 0-100>,
  "score_breakdown": {{
    "formatting": <integer 0-25>,
    "keyword_optimization": <integer 0-25>,
    "content_quality": <integer 0-25>,
    "structure_sections": <integer 0-25>
  }},
  "strengths": ["<short point>", "..."],
  "critical_issues": ["<short point>", "..."],
  "section_feedback": {{
    "contact_info": "<feedback + ATS-friendly fix>",
    "summary_objective": "<feedback + ATS-friendly fix>",
    "work_experience": "<feedback + ATS-friendly fix, mention bullet phrasing/action verbs/quantification>",
    "skills": "<feedback + ATS-friendly fix, keyword suggestions>",
    "education": "<feedback + ATS-friendly fix>",
    "formatting_layout": "<feedback about tables/columns/graphics/fonts that break ATS parsing>"
  }},
  "missing_keywords": ["<keyword>", "..."],
  "rewritten_bullet_examples": [
    {{"original": "<a weak bullet found in resume, or 'N/A' if none found>", "improved": "<stronger ATS-friendly, quantified rewrite>"}}
  ]
}}

Rules for the improvements you suggest:
- Every suggestion must be genuinely ATS-friendly: standard section headings, no tables/text boxes/columns/graphics/icons for critical info, standard fonts, no headers/footers for key content, simple bullet points, spelled-out + abbreviated forms of key terms where relevant (e.g. "Search Engine Optimization (SEO)").
- Favor strong action verbs and quantifiable results (numbers, %, $, time saved) in bullet rewrites.
- Keep feedback specific to THIS resume's actual content, not generic advice.
- If a section is missing entirely, say so and explain what to add.
"""


def build_jd_section(job_description):
    if job_description and job_description.strip():
        return f"""ALSO compare it against this target job description and weight missing_keywords and keyword_optimization accordingly:
\"\"\"
{job_description}
\"\"\""""
    return "No specific job description was provided — evaluate against general ATS best practices for this resume's apparent field."


def analyze_resume(resume_text, job_description, client):
    prompt = ATS_PROMPT_TEMPLATE.format(
        resume_text=resume_text,
        jd_section=build_jd_section(job_description),
    )
    try:
        response = client.models.generate_content(
            model=MODEL_NAME,
            contents=prompt,
            config=types.GenerateContentConfig(temperature=0.3),
        )
    except Exception as e:
        st.error(f"Gemini API request failed: {e}")
        st.info(
            "Common causes: invalid/expired API key, no free quota left today, "
            "or a temporary Gemini API outage. Check your key at "
            "https://aistudio.google.com/app/apikey"
        )
        return None
    raw = response.text.strip()

    # Strip markdown code fences if Gemini adds them despite instructions
    raw = re.sub(r"^```json\s*|^```\s*|```$", "", raw, flags=re.MULTILINE).strip()

    try:
        data = json.loads(raw)
        return data
    except json.JSONDecodeError:
        st.error("Couldn't parse the AI response as JSON. Raw response shown below for debugging.")
        st.code(raw)
        return None


# ---------- UI ----------

st.title("📄 ATS Resume Scorer & Improver")
st.write("Upload your resume (PDF, DOCX, or image) and get an ATS score plus concrete, ATS-friendly improvements.")

col1, col2 = st.columns([2, 1])
with col1:
    uploaded_file = st.file_uploader(
        "Upload your resume", type=["pdf", "docx", "doc", "png", "jpg", "jpeg"]
    )
with col2:
    job_description = st.text_area(
        "Paste target Job Description (optional, improves accuracy)", height=150
    )

analyze_btn = st.button("🔍 Analyze Resume", type="primary")

if analyze_btn:
    if not api_key or client is None:
        st.error("Please enter your Gemini API key in the sidebar first.")
    elif not uploaded_file:
        st.error("Please upload a resume file.")
    else:
        with st.spinner("Extracting text from your resume..."):
            resume_text = get_resume_text(uploaded_file, client)

        if not resume_text:
            st.error("Couldn't extract any text from this file. Try a different format.")
        else:
            with st.expander("📋 Extracted Resume Text (verify this looks correct)"):
                st.text(resume_text)

            with st.spinner("Analyzing with Gemini..."):
                result = analyze_resume(resume_text, job_description, client)

            if result:
                score = result.get("ats_score", 0)
                st.header(f"ATS Score: {score}/100")
                st.progress(min(max(score, 0), 100) / 100)

                breakdown = result.get("score_breakdown", {})
                if breakdown:
                    b1, b2, b3, b4 = st.columns(4)
                    b1.metric("Formatting", f"{breakdown.get('formatting', 0)}/25")
                    b2.metric("Keywords", f"{breakdown.get('keyword_optimization', 0)}/25")
                    b3.metric("Content Quality", f"{breakdown.get('content_quality', 0)}/25")
                    b4.metric("Structure", f"{breakdown.get('structure_sections', 0)}/25")

                st.subheader("✅ Strengths")
                for s in result.get("strengths", []):
                    st.markdown(f"- {s}")

                st.subheader("⚠️ Critical Issues")
                for i in result.get("critical_issues", []):
                    st.markdown(f"- {i}")

                st.subheader("🛠️ Section-by-Section Improvements")
                section_feedback = result.get("section_feedback", {})
                for section, feedback in section_feedback.items():
                    with st.expander(section.replace("_", " ").title()):
                        st.write(feedback)

                missing_kw = result.get("missing_keywords", [])
                if missing_kw:
                    st.subheader("🔑 Missing Keywords to Add")
                    st.write(", ".join(missing_kw))

                rewrites = result.get("rewritten_bullet_examples", [])
                if rewrites:
                    st.subheader("✍️ Bullet Point Rewrites")
                    for r in rewrites:
                        st.markdown(f"**Original:** {r.get('original', '')}")
                        st.markdown(f"**Improved:** {r.get('improved', '')}")
                        st.divider()

                st.download_button(
                    "📥 Download Full Report (JSON)",
                    data=json.dumps(result, indent=2),
                    file_name="ats_report.json",
                    mime="application/json",
                )

st.sidebar.markdown("---")
st.sidebar.caption("Your API key is only used for this session and is never stored.")
